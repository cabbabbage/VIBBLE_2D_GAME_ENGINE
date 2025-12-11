from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

root = Path.cwd()
html_path = root / 'overview.html'
html = html_path.read_text(encoding='utf-8')
soap = BeautifulSoup(html, 'html.parser')

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2

doc.add_heading('CS462 Final Submission', 0)
doc.add_paragraph()

heading_levels = {f'h{i}': i for i in range(1, 7)}

def process_element(el):
    if el.name is None:
        return
    if el.name in heading_levels:
        text = el.get_text(strip=True)
        if text:
            doc.add_heading(text, level=heading_levels[el.name])
        return
    if el.name == 'p':
        text = el.get_text(strip=True)
        if text:
            doc.add_paragraph(text)
        return
    if el.name == 'figure':
        img = el.find('img')
        if img and img.get('src'):
            image_path = root / Path(img['src'])
            if image_path.exists():
                doc.add_picture(str(image_path))
            else:
                doc.add_paragraph(f'Image missing: {img["src"]}')
        caption = el.find('figcaption')
        if caption:
            cap_par = doc.add_paragraph(caption.get_text(strip=True))
            if cap_par.runs:
                cap_par.runs[0].italic = True
        return
    if el.name in ('section', 'article', 'body', 'html'):
        for child in el.children:
            if getattr(child, 'name', None) is None:
                continue
            process_element(child)
        return

body = soap.body
if body:
    process_element(body)

output_path = root / 'CS462 Final Submission.docx'
doc.save(str(output_path))
print('Saved', output_path)
