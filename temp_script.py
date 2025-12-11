from docx import Document
from pathlib import Path
path = Path('CS462 Final Submission.docx')
doc = Document(path)
start=None
end=None
for i,p in enumerate(doc.paragraphs):
    if p.text.strip()=="Technical Details":
        start=i
    if p.text.strip()=="Game Challenge" and start is not None and end is None:
        end=i
        break
if start is None or end is None:
    raise SystemExit('bounds missing')
for i in range(start, end):
    p=doc.paragraphs[i]
    if p.text.strip():
        print(i, repr(p.text))
